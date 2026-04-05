"""
Stage 1 — Document Ingestion & Text Extraction
Identifies file type and extracts text using appropriate strategy:
  PDF → PyMuPDF page-by-page
  DOCX → python-docx paragraphs + tables
  TXT/MD → direct read
"""
import json
from pathlib import Path
from typing import List, Dict

from backend.config import CLIENTS_DIR
from backend.services.audit_service import emit_audit_event


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from PDF page by page using PyMuPDF."""
    import fitz  # PyMuPDF
    doc = fitz.open(str(file_path))
    pages = []
    for page_num, page in enumerate(doc, 1):
        text = page.get_text()
        if text.strip():
            pages.append(f"--- Page {page_num} ---\n{text}")
    doc.close()
    return "\n\n".join(pages)


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from DOCX preserving paragraphs and tables."""
    from docx import Document
    doc = Document(str(file_path))
    parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if text:
                        # Preserve heading hierarchy
                        if para.style and para.style.name.startswith("Heading"):
                            level = para.style.name.replace("Heading ", "").strip()
                            parts.append(f"{'#' * int(level) if level.isdigit() else '##'} {text}")
                        else:
                            parts.append(text)
                    break

        elif tag == "tbl":
            # Table
            for table in doc.tables:
                if table._element == element:
                    table_text = []
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        table_text.append(row_text)
                    if table_text:
                        parts.append("\n".join(table_text))
                    break

    return "\n\n".join(parts)


def extract_text_from_txt(file_path: Path) -> str:
    """Read plain text or markdown file directly."""
    return file_path.read_text(encoding="utf-8")


def run_stage1(client_id: str) -> Dict[str, str]:
    """
    Execute Stage 1 — Document Ingestion.
    
    Scans /input_documents/ for all uploaded files, extracts text from each.
    
    Args:
        client_id: The client folder ID
        
    Returns:
        Dict mapping filename → extracted text
    """
    print(f"\n{'='*60}")
    print(f"  Stage 1 — Document Ingestion & Text Extraction")
    print(f"{'='*60}")

    docs_dir = CLIENTS_DIR / client_id / "input_documents"
    if not docs_dir.exists() or not any(docs_dir.iterdir()):
        raise FileNotFoundError(f"No documents found in {docs_dir}")

    extracted = {}
    supported_extensions = {".pdf", ".docx", ".txt", ".md", ".markdown"}

    for file_path in sorted(docs_dir.iterdir()):
        if not file_path.is_file():
            continue
        ext = file_path.suffix.lower()
        if ext not in supported_extensions:
            print(f"  ⚠️  Skipping unsupported file: {file_path.name}")
            continue

        print(f"  📄 Extracting: {file_path.name} ({ext})")

        try:
            if ext == ".pdf":
                text = extract_text_from_pdf(file_path)
            elif ext == ".docx":
                text = extract_text_from_docx(file_path)
            else:
                text = extract_text_from_txt(file_path)

            extracted[file_path.name] = text
            print(f"     ✅ Extracted {len(text)} characters")

        except Exception as e:
            print(f"     ❌ Failed to extract {file_path.name}: {e}")
            extracted[file_path.name] = f"[EXTRACTION_FAILED: {str(e)}]"

    # Audit
    total_chars = sum(len(t) for t in extracted.values())
    emit_audit_event(
        client_id=client_id,
        stage="stage_1_ingestion",
        action=f"Extracted text from {len(extracted)} documents ({total_chars} chars total)",
        agent="stage1_ingestion",
        input_data=str(list(extracted.keys())),
        output_data=str(total_chars),
    )

    print(f"\n  ✅ Stage 1 complete — {len(extracted)} documents processed")
    return extracted
